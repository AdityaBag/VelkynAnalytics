from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
import base64
import json
import re
import os

from docx import Document
from docx.shared import Inches
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

from backend.report_training import (
    SUPPORTED_TRAINING_EXTS,
    build_style_profile,
    extract_text_from_file,
    load_style_profile,
    save_style_profile,
    slugify,
)


router = APIRouter(prefix="/report-agent", tags=["Report Agent"])

IS_LAMBDA = bool(os.getenv("AWS_LAMBDA_FUNCTION_NAME"))
BASE_WRITE_ROOT = Path("/tmp") if IS_LAMBDA else Path(".")

TRAINING_ROOT = BASE_WRITE_ROOT / "training_materials"
TRAINING_UPLOADS = TRAINING_ROOT / "uploads"
TRAINING_PROFILES = TRAINING_ROOT / "profiles"
REPORT_OUTPUT_ROOT = BASE_WRITE_ROOT / "results" / "reports"

for folder in (TRAINING_ROOT, TRAINING_UPLOADS, TRAINING_PROFILES, REPORT_OUTPUT_ROOT):
    folder.mkdir(parents=True, exist_ok=True)

REPORT_TYPES = [
    "Pricing & Valuation Report",
    "Greeks Exposure Report",
    "Scenario Shock Report",
    "Multi-Rate Sensitivity Report",
    "Volatility Smile Report",
    "Volatility Surface Report",
    "Batch Pricing Report",
    "Run History / Audit Report",
    "Monte Carlo Diagnostics Report",
    "Binomial Convergence Report",
    "Model Comparison Report",
    "Live Market Snapshot Report",
]

REPORT_STYLES = [
    "Model Validation Style",
    "Sell-Side Quant Research Style",
    "Academic Quant Style",
    "Risk & Stress Testing Style",
    "Concise Technical Summary",
    "Verbose Research Report",
]


class ChartImage(BaseModel):
    name: str = Field(default="chart")
    data_url: str = Field(...)


class GenerateReportRequest(BaseModel):
    report_type: str
    report_style: str
    author: str | None = None
    dash_shared: dict[str, Any] = Field(default_factory=dict)
    snapshots: dict[str, Any] = Field(default_factory=dict)
    charts: list[ChartImage] = Field(default_factory=list)
    notes: str | None = None


def _profile_path(report_type: str) -> Path:
    return TRAINING_PROFILES / f"{slugify(report_type)}.json"


def _load_report_profile(report_type: str) -> dict[str, Any] | None:
    return load_style_profile(_profile_path(report_type))


def _default_status(report_type: str) -> dict[str, Any]:
    return {
        "report_type": report_type,
        "enabled": False,
        "reason": "Upload training files to enable this report.",
        "capabilities": {
            "structure": False,
            "tone": False,
            "narrative": False,
            "examples": False,
            "ready": False,
        },
    }


def _report_type_status(report_type: str) -> dict[str, Any]:
    profile = _load_report_profile(report_type)
    if not profile:
        return _default_status(report_type)

    caps = profile.get("capabilities", {})
    ready = bool(caps.get("ready", False))
    reason = "Ready" if ready else "Insufficient training signal (structure/tone/narrative/examples)."
    return {
        "report_type": report_type,
        "enabled": ready,
        "reason": reason,
        "capabilities": {
            "structure": bool(caps.get("structure", False)),
            "tone": bool(caps.get("tone", False)),
            "narrative": bool(caps.get("narrative", False)),
            "examples": bool(caps.get("examples", False)),
            "ready": ready,
        },
        "updated_at": profile.get("updated_at"),
        "doc_count": profile.get("doc_count", 0),
    }


def _style_voice(style: str, profile: dict[str, Any]) -> str:
    tone = (profile or {}).get("tone", {}).get("dominant", "formal")
    if style == "Model Validation Style":
        return "validation-oriented, methodical, and control-focused"
    if style == "Sell-Side Quant Research Style":
        return "market-facing, thesis-driven, and catalyst-aware"
    if style == "Academic Quant Style":
        return "rigorous, citation-ready, and hypothesis-centered"
    if style == "Risk & Stress Testing Style":
        return "risk-first, stress-centric, and downside-aware"
    if style == "Concise Technical Summary":
        return "compact, technical, and action-oriented"
    if style == "Verbose Research Report":
        return "expansive, narrative-rich, and deeply analytical"
    return f"{tone} and technically precise"


def _flatten(prefix: str, value: Any, out: dict[str, Any]) -> None:
    if isinstance(value, dict):
        for key, item in value.items():
            next_prefix = f"{prefix}.{key}" if prefix else str(key)
            _flatten(next_prefix, item, out)
        return

    if isinstance(value, list):
        if not value:
            out[prefix] = "[]"
            return
        if all(not isinstance(item, (dict, list)) for item in value):
            out[prefix] = ", ".join(str(item) for item in value[:20])
        else:
            out[prefix] = f"list[{len(value)}]"
        return

    out[prefix] = value


def _add_key_value_table(doc: Document, title: str, payload: dict[str, Any], max_rows: int = 24) -> None:
    if not payload:
        return

    flat: dict[str, Any] = {}
    _flatten("", payload, flat)
    rows = list(flat.items())[:max_rows]
    if not rows:
        return

    doc.add_heading(title, level=2)
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Light List Accent 1"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    for idx, (key, value) in enumerate(rows, start=1):
        table.rows[idx].cells[0].text = str(key)
        table.rows[idx].cells[1].text = str(value)


def _add_results_tables(doc: Document, dash_shared: dict[str, Any], snapshots: dict[str, Any]) -> None:
    doc.add_heading("Results", level=1)

    sections = []
    if isinstance(dash_shared, dict):
        sections.extend((f"Dash Shared: {k}", v) for k, v in dash_shared.items() if isinstance(v, dict))
    if isinstance(snapshots, dict):
        sections.extend((f"Snapshot: {k}", v) for k, v in snapshots.items() if isinstance(v, dict))

    emitted = 0
    for title, payload in sections:
        if emitted >= 8:
            break
        _add_key_value_table(doc, title, payload, max_rows=18)
        emitted += 1

    if emitted == 0:
        doc.add_paragraph("No structured result snapshots were supplied by the frontend payload.")


def _decode_chart_data_url(data_url: str) -> bytes | None:
    if not isinstance(data_url, str) or not data_url.startswith("data:image"):
        return None
    match = re.match(r"^data:image\/[a-zA-Z0-9.+-]+;base64,(.+)$", data_url)
    if not match:
        return None
    try:
        return base64.b64decode(match.group(1))
    except Exception:
        return None


def _add_charts(doc: Document, charts: list[ChartImage]) -> None:
    doc.add_heading("Results Charts", level=2)
    added = 0
    for chart in charts[:12]:
        raw = _decode_chart_data_url(chart.data_url)
        if not raw:
            continue

        doc.add_paragraph(chart.name or f"Chart {added + 1}")
        stream = BytesIO(raw)
        try:
            doc.add_picture(stream, width=Inches(6.2))
            added += 1
        except Exception:
            continue

    if added == 0:
        doc.add_paragraph("No chart images were provided in the payload.")


def _compose_section_text(section_name: str, req: GenerateReportRequest, profile: dict[str, Any]) -> str:
    phrases = profile.get("common_phrases", [])[:3]
    phrase_text = "; ".join(phrases) if phrases else "model consistency, parameter transparency, and reproducibility"
    voice = _style_voice(req.report_style, profile)

    snippets = {
        "Executive summary": f"This {req.report_type} is prepared in a {voice} register. The report synthesizes current engine outputs and dashboard snapshots, emphasizing {phrase_text}.",
        "Methodology": "The workflow aggregates engine-level payloads, harmonizes assumptions, and applies cross-checks between primary pricing, risk, and scenario blocks before producing narrative conclusions.",
        "Interpretation and commentary": "Observed outputs are interpreted relative to baseline assumptions and scenario context. Where divergence appears across models, commentary highlights likely drivers and follow-up checks.",
        "Limitations": "Results remain dependent on input quality, discretization choices, and available market snapshots. Report statements should be validated against live data and independent controls before external use.",
        "Next steps": "Recommended actions include rerunning stressed cases, validating sensitivity stability, and extending training materials for tighter style alignment in future report generations.",
    }
    return snippets.get(section_name, "")


@router.get("/meta")
def report_agent_meta():
    statuses = [_report_type_status(item) for item in REPORT_TYPES]
    return {
        "report_types": REPORT_TYPES,
        "report_styles": REPORT_STYLES,
        "status": statuses,
        "supported_training_formats": sorted(SUPPORTED_TRAINING_EXTS),
    }


@router.post("/train")
async def train_report_profile(report_type: str = Form(...), files: list[UploadFile] = File(...)):
    if report_type not in REPORT_TYPES:
        raise HTTPException(status_code=400, detail="Unknown report type")
    if not files:
        raise HTTPException(status_code=400, detail="At least one training file is required")

    report_slug = slugify(report_type)
    target_dir = TRAINING_UPLOADS / report_slug
    target_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    for file in files:
        suffix = Path(file.filename or "").suffix.lower()
        if suffix not in SUPPORTED_TRAINING_EXTS:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {suffix}")

        safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", Path(file.filename or f"upload{suffix}").name)
        out_path = target_dir / f"{ts}_{safe_name}"
        data = await file.read()
        out_path.write_bytes(data)

    docs: list[dict[str, Any]] = []
    for existing_file in sorted(target_dir.iterdir(), key=lambda p: p.name.lower()):
        if not existing_file.is_file() or existing_file.suffix.lower() not in SUPPORTED_TRAINING_EXTS:
            continue
        try:
            text = extract_text_from_file(existing_file)
        except Exception:
            continue
        stat = existing_file.stat()
        docs.append(
            {
                "name": existing_file.name,
                "path": str(existing_file).replace("\\", "/"),
                "size": int(stat.st_size),
                "chars": len(text or ""),
                "text": text or "",
                "uploaded_at": datetime.fromtimestamp(stat.st_mtime).isoformat() + "Z",
            }
        )

    profile = build_style_profile(report_type, docs)
    profile_path = _profile_path(report_type)

    save_style_profile(profile_path, profile)

    status = _report_type_status(report_type)
    return {
        "ok": True,
        "report_type": report_type,
        "profile_path": str(profile_path).replace("\\", "/"),
        "status": status,
    }


@router.post("/generate")
def generate_word_report(req: GenerateReportRequest):
    if req.report_type not in REPORT_TYPES:
        raise HTTPException(status_code=400, detail="Unknown report type")
    if req.report_style not in REPORT_STYLES:
        raise HTTPException(status_code=400, detail="Unknown report style")

    profile = _load_report_profile(req.report_type)
    if not profile or not bool(profile.get("capabilities", {}).get("ready", False)):
        raise HTTPException(
            status_code=400,
            detail="Selected report type is not trained enough yet. Upload richer training files first.",
        )

    document = Document()
    document.core_properties.author = req.author or "QPS Report Agent"
    document.core_properties.title = req.report_type

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    document.add_heading(req.report_type, level=0)
    document.add_paragraph(f"Date: {generated_at}")
    document.add_paragraph(f"Author: {req.author or 'QPS Report Agent'}")
    document.add_paragraph(f"Style: {req.report_style}")
    document.add_page_break()

    document.add_heading("Executive summary", level=1)
    document.add_paragraph(_compose_section_text("Executive summary", req, profile))

    document.add_heading("Inputs and assumptions", level=1)
    assumptions = {
        "report_type": req.report_type,
        "report_style": req.report_style,
        "training_profile_updated": profile.get("updated_at"),
        "training_doc_count": profile.get("doc_count", 0),
        "snapshot_keys": sorted((req.snapshots or {}).keys()),
        "dash_shared_keys": sorted((req.dash_shared or {}).keys()),
    }
    if req.notes:
        assumptions["user_notes"] = req.notes
    _add_key_value_table(document, "Input Snapshot", assumptions, max_rows=20)

    document.add_heading("Methodology", level=1)
    document.add_paragraph(_compose_section_text("Methodology", req, profile))

    _add_results_tables(document, req.dash_shared or {}, req.snapshots or {})
    _add_charts(document, req.charts or [])

    document.add_heading("Interpretation and commentary", level=1)
    document.add_paragraph(_compose_section_text("Interpretation and commentary", req, profile))

    document.add_heading("Limitations", level=1)
    document.add_paragraph(_compose_section_text("Limitations", req, profile))

    document.add_heading("Next steps", level=1)
    document.add_paragraph(_compose_section_text("Next steps", req, profile))

    out_day = datetime.now().strftime("%Y-%m-%d")
    out_dir = REPORT_OUTPUT_ROOT / out_day
    out_dir.mkdir(parents=True, exist_ok=True)

    file_stem = f"{datetime.now().strftime('%H-%M-%S')}_{slugify(req.report_type)}"
    output_path = out_dir / f"{file_stem}.docx"
    document.save(str(output_path))

    filename = output_path.name
    return FileResponse(
        path=output_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )
